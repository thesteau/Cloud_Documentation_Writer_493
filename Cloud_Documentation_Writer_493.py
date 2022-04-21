from docx import Document
from docx.shared import Pt


class PythonAPIWriter:
    """ Generate an API document per specifications."""

    def __init__(self, doc_name='word'):
        self.document = Document()
        self.doc_name = doc_name + '.docx'

    def write_page(self, context=None):
        self.page_end()
        self.write_head(context)

        self.write_request(context)
        self.write_response(context)

        self.page_end()

    def write_head(self, context):
        # Head name
        if context is not None:
            self.write_heading(context["head"], "Heading 1", 16)

            # Head subtext
            self.write_text(context["head_details"])

            # Add route
            self.write_code(context["head_route"])
        else:
            self.write_heading("Method", "Heading 1", 16)

            # Head subtext
            self.write_text("Something")

            # Add route
            self.write_code("ROUTE")

    def write_request(self, context):
        # Request
        self.write_heading("Request", "Heading 2", 13)

        self.write_heading("Path Parameters", "Heading 3", 12)
        self.write_table(['Name', 'Description'])

        self.write_heading("Request Body", "Heading 3", 12)

        self.write_condition("None", context, "req_body")

        self.write_heading("Request Body Format", "Heading 3", 12)
        self.write_condition("None", context, "req_body_format")

        self.write_heading("Request JSON Attributes", "Heading 3", 12)
        self.write_table(['Name', 'Description', 'Required?'])

        self.write_heading("Request Body Example", "Heading 3", 12)
        self.write_code()

    def write_response(self, context):
        self.write_heading("Response", "Heading 2", 12)

        self.write_heading("Response Body Format", "Heading 3", 12)
        self.write_condition("None", context, "res_body_format")

        self.write_heading("Response Statuses", "Heading 3", 12)
        self.write_table(['Outcome', 'Status Code', 'Notes'])

        self.write_heading("Response Examples", "Heading 3", 12)

        self.write_heading("Success", "Heading 4", 11)
        self.write_code()

        self.write_heading("Failure", "Heading 4", 11)
        self.write_code()

        self.write_heading("Notes", "Heading 3", 12)

    def write_heading(self, text, style, size):
        heading = self.document.add_paragraph()
        heading.text = text
        heading.style = self.document.styles[style]
        heading.style.font.name = "Calibri Light (Headings)"
        heading.style.font.size = Pt(size)
        heading.style.font.bold = False

    def write_text(self, text):
        text_data = self.document.add_paragraph()
        text_data.style = self.document.styles["Normal"]
        text_data.text = text

    def write_table(self, text_list):
        the_cols = len(text_list)
        table = self.document.add_table(rows=2, cols=the_cols)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells

        for each_cell in range(the_cols):
            cell = hdr_cells[each_cell].paragraphs[0]
            cell.add_run(text_list[each_cell]).bold = True

    def write_code(self, text=None):
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'TableGrid'
        if text is None:
            table.cell(0, 0).text = ""
        else:
            table.cell(0, 0).text = text

    def write_condition(self, main, context, call=None):
        if context is None:
            self.write_text(main)
        else:
            self.write_text(context[call])

    def page_end(self):
        self.document.add_page_break()

    def save_doc(self):
        self.document.save(self.doc_name)


if __name__ == '__main__':

    context = [
        {
            "head": "Create a Boat",
            "head_details": "Allows you to create a new boat.",
            "head_route": "POST /boats",
            "req_body": 'Required',
            "req_body_format": 'JSON',
            "res_body_format": "JSON"
        }
    ]

    api_write = PythonAPIWriter()
    api_write.write_page(context[0])
    api_write.save_doc()
